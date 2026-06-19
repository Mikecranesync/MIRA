"""Heavy document extraction → a normalized Document IR.

Reads ANY document (PDF digital + scanned, Word, Excel, images, text/markdown/html, csv) into a flat
list of ``DocBlock``s that the deterministic contextualization layer (P2) consumes. OCR is built in
(Tesseract via pytesseract; scanned-PDF pages rasterized with pypdfium2) and degrades gracefully when
the Tesseract binary is absent — extraction never crashes, it records a warning and a low-confidence
block.

All third-party imports are LAZY (inside the per-format functions) so importing this module — and the
P0 core — needs none of the heavy deps. Install them with the ``[docs]`` extra.
"""
from __future__ import annotations

import os
import re
from dataclasses import dataclass, field

# Confidence bands match docs/specs (uns-message-resolver §2.4): high/medium/low.
_HIGH, _MED, _LOW = "high", "medium", "low"

TEXT_EXTS = {".txt", ".md", ".log"}
HTML_EXTS = {".html", ".htm"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif"}
DOC_EXTS = TEXT_EXTS | HTML_EXTS | IMAGE_EXTS | {".pdf", ".docx", ".xlsx", ".csv"}


@dataclass
class DocBlock:
    text: str
    kind: str = "text"            # text | table | ocr
    page: int | None = None
    section: str | None = None
    confidence: str = _HIGH       # extraction confidence (not contextualization confidence)

    def to_dict(self) -> dict:
        return {"text": self.text, "kind": self.kind, "page": self.page,
                "section": self.section, "confidence": self.confidence}


@dataclass
class ExtractResult:
    file_name: str
    extractor: str
    blocks: list[DocBlock] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n".join(b.text for b in self.blocks if b.text).strip()

    def to_dict(self) -> dict:
        return {"schema": "mira-contextualizer/document@1", "file_name": self.file_name,
                "extractor": self.extractor, "blocks": [b.to_dict() for b in self.blocks],
                "warnings": list(self.warnings)}


def can_extract(file_name: str) -> bool:
    return os.path.splitext(file_name)[1].lower() in DOC_EXTS


def ocr_available() -> bool:
    """True if the Tesseract binary is reachable (pytesseract installed AND engine on PATH)."""
    try:
        import pytesseract  # noqa: PLC0415
        pytesseract.get_tesseract_version()
        return True
    except Exception:  # noqa: BLE001 — ImportError or TesseractNotFoundError
        return False


def _ocr_image(img) -> tuple[str, bool]:
    """OCR a PIL image. Returns (text, ok); ok=False when the Tesseract engine is unavailable."""
    try:
        import pytesseract  # noqa: PLC0415
        return pytesseract.image_to_string(img) or "", True
    except Exception:  # noqa: BLE001
        return "", False


def extract(path: str, file_name: str | None = None) -> ExtractResult:
    """Route a file to its extractor by extension. Always returns a result (warnings on failure)."""
    file_name = file_name or os.path.basename(path)
    ext = os.path.splitext(file_name)[1].lower()
    try:
        if ext in TEXT_EXTS:
            return _extract_text(path, file_name)
        if ext in HTML_EXTS:
            return _extract_html(path, file_name)
        if ext == ".csv":
            return _extract_csv(path, file_name)
        if ext == ".xlsx":
            return _extract_xlsx(path, file_name)
        if ext == ".docx":
            return _extract_docx(path, file_name)
        if ext == ".pdf":
            return _extract_pdf(path, file_name)
        if ext in IMAGE_EXTS:
            return _extract_image(path, file_name)
    except Exception as exc:  # noqa: BLE001 — never crash the caller on a bad file
        return ExtractResult(file_name, "error", warnings=["extraction failed: %s" % exc])
    return ExtractResult(file_name, "none", warnings=["unsupported document type: %s" % ext])


# ── per-format extractors ────────────────────────────────────────────────────
def _extract_text(path: str, file_name: str) -> ExtractResult:
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return ExtractResult(file_name, "text", [DocBlock(fh.read())])


def _extract_html(path: str, file_name: str) -> ExtractResult:
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        raw = fh.read()
    raw = re.sub(r"(?is)<(script|style)\b.*?</\1>", " ", raw)
    text = re.sub(r"(?s)<[^>]+>", " ", raw)
    text = re.sub(r"[ \t]+", " ", re.sub(r"&nbsp;", " ", text)).strip()
    return ExtractResult(file_name, "html", [DocBlock(text)])


def _extract_csv(path: str, file_name: str) -> ExtractResult:
    import csv  # noqa: PLC0415
    with open(path, "r", encoding="utf-8", errors="replace", newline="") as fh:
        rows = list(csv.reader(fh))
    body = "\n".join("\t".join(r) for r in rows)
    return ExtractResult(file_name, "csv", [DocBlock(body, kind="table")])


def _extract_xlsx(path: str, file_name: str) -> ExtractResult:
    from openpyxl import load_workbook  # noqa: PLC0415
    wb = load_workbook(path, read_only=True, data_only=True)
    blocks: list[DocBlock] = []
    for ws in wb.worksheets:
        rows = ["\t".join("" if c is None else str(c) for c in row) for row in ws.iter_rows(values_only=True)]
        body = "\n".join(rows).strip()
        if body:
            blocks.append(DocBlock(body, kind="table", section=ws.title))
    wb.close()
    return ExtractResult(file_name, "xlsx", blocks or [DocBlock("", confidence=_LOW)])


def _extract_docx(path: str, file_name: str) -> ExtractResult:
    import docx  # noqa: PLC0415
    doc = docx.Document(path)
    blocks = [DocBlock(p.text) for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        rows = ["\t".join(c.text for c in row.cells) for row in t.rows]
        body = "\n".join(rows).strip()
        if body:
            blocks.append(DocBlock(body, kind="table"))
    return ExtractResult(file_name, "docx", blocks or [DocBlock("", confidence=_LOW)])


def _extract_pdf(path: str, file_name: str) -> ExtractResult:
    from pdfminer.high_level import extract_text  # noqa: PLC0415

    try:
        import pypdfium2 as pdfium  # noqa: PLC0415
        n_pages = len(pdfium.PdfDocument(path))
    except Exception:  # noqa: BLE001 — fall back to single-shot text if rasterizer is unavailable
        pdfium, n_pages = None, 0

    res = ExtractResult(file_name, "pdf")
    if not n_pages:
        text = (extract_text(path) or "").strip()
        res.blocks.append(DocBlock(text, page=None, confidence=_HIGH if text else _LOW))
        if not text:
            res.warnings.append("no embedded text and page rasterizer unavailable")
        return res

    pdf = pdfium.PdfDocument(path)
    ocr_ok_seen = False
    for i in range(n_pages):
        text = (extract_text(path, page_numbers=[i]) or "").strip()
        if len(text) >= 10:
            res.blocks.append(DocBlock(text, page=i + 1))
            continue
        # likely a scanned / image-only page → rasterize and OCR
        img = pdf[i].render(scale=2.0).to_pil()
        ocr_text, ok = _ocr_image(img)
        ocr_ok_seen = ocr_ok_seen or ok
        res.blocks.append(DocBlock(
            ocr_text.strip(), kind="ocr", page=i + 1,
            confidence=_MED if ocr_text.strip() else _LOW,
        ))
    if not ocr_ok_seen and any(b.kind == "ocr" for b in res.blocks):
        res.warnings.append("scanned pages detected but Tesseract OCR engine unavailable")
    return res


def _extract_image(path: str, file_name: str) -> ExtractResult:
    from PIL import Image  # noqa: PLC0415
    with Image.open(path) as img:
        text, ok = _ocr_image(img.convert("RGB"))
    res = ExtractResult(file_name, "image-ocr",
                        [DocBlock(text.strip(), kind="ocr", confidence=_MED if text.strip() else _LOW)])
    if not ok:
        res.warnings.append("Tesseract OCR engine unavailable")
    return res
