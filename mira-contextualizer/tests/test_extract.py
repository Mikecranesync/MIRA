"""Heavy document extraction → Document IR. Real extraction for digital formats; OCR degrades
gracefully when the Tesseract engine is absent (the dev/CI case)."""
import importlib.util

import pytest

from mira_contextualizer import extract

_HAS = lambda m: importlib.util.find_spec(m) is not None  # noqa: E731


def test_can_extract_and_ocr_available_is_bool():
    assert extract.can_extract("manual.pdf") and extract.can_extract("a.PNG")
    assert not extract.can_extract("program.l5x")  # PLC routes to the parser, not here
    assert isinstance(extract.ocr_available(), bool)


def test_text_and_html(tmp_path):
    p = tmp_path / "n.txt"; p.write_text("Overload on CV-101\nreset breaker", encoding="utf-8")
    r = extract.extract(str(p))
    assert r.extractor == "text" and "Overload on CV-101" in r.full_text

    h = tmp_path / "n.html"
    h.write_text("<html><head><style>x{}</style></head><body><h1>VFD</h1> fault F0004</body></html>", encoding="utf-8")
    rh = extract.extract(str(h))
    assert "VFD" in rh.full_text and "fault F0004" in rh.full_text and "<" not in rh.full_text


def test_csv_table(tmp_path):
    p = tmp_path / "bom.csv"; p.write_text("part,qty\nGS10,1\nCV-101,2\n", encoding="utf-8")
    r = extract.extract(str(p))
    assert r.extractor == "csv" and r.blocks[0].kind == "table" and "GS10" in r.full_text


@pytest.mark.skipif(not _HAS("openpyxl"), reason="openpyxl not installed")
def test_xlsx(tmp_path):
    from openpyxl import Workbook
    wb = Workbook(); ws = wb.active; ws.title = "Params"
    ws.append(["param", "value"]); ws.append(["P09.03", "5s"])
    p = tmp_path / "p.xlsx"; wb.save(str(p))
    r = extract.extract(str(p))
    assert r.extractor == "xlsx" and "P09.03" in r.full_text
    assert any(b.section == "Params" for b in r.blocks)


@pytest.mark.skipif(not _HAS("docx"), reason="python-docx not installed")
def test_docx(tmp_path):
    import docx
    d = docx.Document(); d.add_paragraph("Drive trips on overload F0004.")
    t = d.add_table(rows=1, cols=2); t.rows[0].cells[0].text = "Code"; t.rows[0].cells[1].text = "Meaning"
    p = tmp_path / "m.docx"; d.save(str(p))
    r = extract.extract(str(p))
    assert r.extractor == "docx" and "F0004" in r.full_text
    assert any(b.kind == "table" for b in r.blocks)


def _min_pdf(text: str) -> bytes:
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        None,  # filled below (stream)
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream = b"BT /F1 14 Tf 20 120 Td (" + text.encode("latin-1") + b") Tj ET"
    objs[3] = b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += str(i).encode() + b" 0 obj\n" + body + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n0 " + str(len(objs) + 1).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        out += ("%010d 00000 n \n" % off).encode()
    out += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() + b" /Root 1 0 R >>\n"
            b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF")
    return bytes(out)


@pytest.mark.skipif(not _HAS("pdfminer"), reason="pdfminer.six not installed")
def test_pdf_digital_text(tmp_path):
    p = tmp_path / "d.pdf"; p.write_bytes(_min_pdf("Fault F0004 overload"))
    r = extract.extract(str(p))
    assert r.extractor == "pdf"
    assert "F0004" in r.full_text


@pytest.mark.skipif(not _HAS("PIL"), reason="Pillow not installed")
def test_image_ocr_graceful(tmp_path):
    from PIL import Image
    img = Image.new("RGB", (60, 30), "white")
    p = tmp_path / "scan.png"; img.save(str(p))
    r = extract.extract(str(p))  # must not crash whether or not Tesseract is installed
    assert r.extractor == "image-ocr" and r.blocks
    if not extract.ocr_available():
        assert any("Tesseract" in w for w in r.warnings)
